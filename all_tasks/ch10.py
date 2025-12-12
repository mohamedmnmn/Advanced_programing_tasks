# import time
# class Timer:
#     def __enter__(self):
#         self.start = time.time()
#         return self
   
#     def __exit__(self, exc_type, exc_val, exc_tb):
#         self.end = time.time()
#         elapsed = self.end - self.start
#         print(f"Execution took {elapsed:.2f} seconds")
# with Timer():
#     for i in range(1000000):
#         pass
# # ---------------------------------------
# def even_numbers(n):
#     for i in range(2, n + 1, 2):
#         yield i
# for num in even_numbers(10):
#     print(num)
# ---------------------------------------
# def filter_positive():
#     while True:
#         num = yield
#         if num > 0:
#             print(f"Positive number: {num}")
# co = filter_positive()
# next(co)
# co.send(-3)
# co.send(5)
# co.send(0)
# --------------------------------------
# class Circle:
#     def draw(self):
#         print("Drawing a Circle")
# class Square:
#     def draw(self):
#         print("Drawing a Square")
# def shape_factory(shape_type):
#     if shape_type == "Circle":
#         return Circle()
#     elif shape_type == "Square":
#         return Square()
# shape = shape_factory("Circle")
# shape.draw()
# shape2 = shape_factory("Square")
# shape2.draw() 
# --------------------------------------
# class Subject:
#     def __init__(self):
#         self.observers = []
#     def attach(self, observer_func):
#         self.observers.append(observer_func)
#     def notify(self, message):
#         for observer in self.observers:
#             observer(message)
# subject = Subject()
# subject.attach(lambda msg: print(f"Observer 1: {msg}"))
# subject.attach(lambda msg: print(f"Observer 2: {msg}"))
# subject.notify("Update available!")
# -------------------------------------------
from docx import Document

files = [r"C:\Users\modye\Desktop\Chapter 1.docx", r"C:\Users\modye\Desktop\chapter2.docx",r"C:\Users\modye\Desktop\chapter3.docx",r"C:\Users\modye\Desktop\chapter4.docx",r"C:\Users\modye\Desktop\chapter5.docx",r"C:\Users\modye\Desktop\chapter6.docx",r"C:\Users\modye\Desktop\chapter7.docx",r"C:\Users\modye\Desktop\chapter8.docx",r"C:\Users\modye\Desktop\chapter9.docx",r"C:\Users\modye\Desktop\chapter10.docx", ]  # your docs
combined = Document()

for f in files:
    doc = Document(f)
    for para in doc.paragraphs:
        combined.add_paragraph(para.text)
    combined.add_page_break()

combined.save("combined_document.docx")
